import sys
import json
from docx import Document
from docx.oxml.ns import qn

def read_docx(file_path):
    doc = Document(file_path)
    result = {
        "paragraphs": [],
        "tables": [],
        "full_text": ""
    }

    # 讀取段落（含標題層級）
    all_text = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name  # e.g. 'Heading 1', 'Normal'
        result["paragraphs"].append({
            "style": style,
            "text": text
        })
        all_text.append(text)

    # 讀取表格
    for table_idx, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        result["tables"].append({
            "table_index": table_idx + 1,
            "data": table_data
        })

    # 合併全文
    result["full_text"] = '\n'.join(all_text)

    # 輸出 JSON 格式，方便 n8n 後續節點使用
    print(json.dumps(result, ensure_ascii=False, indent=2))

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"error": "請提供 docx 檔案路徑"}, ensure_ascii=False))
        sys.exit(1)
    read_docx(sys.argv[1])
